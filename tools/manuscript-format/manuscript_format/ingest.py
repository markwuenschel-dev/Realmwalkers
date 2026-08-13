"""Input side of the manuscript formatter: a file on disk → a :class:`~.spine.Manuscript`.

This module is the INVERSE of ``renderMarkdown`` in ``frontend/src/desk/lib/docx.ts``. Three
recovery paths, in descending fidelity:

1. **Semantic Markdown** (``parse_semantic_markdown``) — a document this tool emitted, carrying
   ``<!-- volume/part/chapter/scene -->`` markers plus a ``schema: dominion-manuscript/v1`` front
   matter block. Lossless for everything ``renderMarkdown`` writes: prose is recovered byte-for-byte
   and structure/labels are recovered from the marker attributes, never re-derived from headings.
2. **Plain Markdown** (``split_plain_markdown``) — ordinary prose with no markers. Structure is
   *inferred* from ``# `` headings and ``***`` scene breaks. Best-effort, not lossless.
3. **DOCX** (``read_docx``) — text recovery only, via ``python-docx``.

Known lossy boundaries (stated up front, because pretending otherwise costs a round-trip):

* **DOCX is a one-way door for panel semantics.** Once a LitRPG ``@interface`` directive has been
  rendered into a Word table by the emitter, the directive that produced it is *gone* — the .docx
  holds a grid of cells, not a role/domain/creature spec. ``read_docx`` therefore recovers prose text
  and table grids, never panel semantics; a round-tripped interface panel comes back as a plain GFM
  pipe table. Character formatting (bold/italic runs) is likewise dropped: plain text is the
  contract. Merged cells repeat their text across the span, and a table nested inside a cell is
  flattened to its paragraph text.
* **``renderMarkdown`` itself drops fields.** ``epigraph`` and ``Manuscript.book_id`` are never
  written (see ``chapterComment``, docx.ts:1584-1590), so no reader can recover them. Scenes with no
  prose and chapters whose scenes are all empty are skipped by ``markdownChapter``
  (docx.ts:1595-1596) and cannot come back.
* **A level-1 heading inside prose is unrecoverable as prose.** ``renderMarkdown`` uses a bare
  ``# `` line at column 0 as its structural delimiter with no escaping, so a scene whose prose
  contains one is ambiguous by construction. This parser resolves the ambiguity by ending the scene,
  matching what any other reader of the file would do. Fenced code blocks ARE protected: ``# `` and
  ``<!-- … -->`` lines inside a ``` fence stay with the prose.
* Prose is recovered with leading/trailing blank lines stripped (the emitter surrounds every scene
  with blank lines and cannot tell them apart from authored ones). Internal blank lines and all
  indentation survive byte-for-byte.
"""

from __future__ import annotations

import re
from collections.abc import Iterator
from dataclasses import dataclass
from dataclasses import field as dc_field
from pathlib import Path

from docx import Document as _open_document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from .labels import SECTION_TYPES
from .spine import (
    Manuscript,
    ManuscriptChapter,
    ManuscriptPart,
    ManuscriptScene,
    ManuscriptVolume,
)

#: The ``schema:`` value ``renderMarkdown`` stamps into front matter (docx.ts:1620).
SEMANTIC_SCHEMA = "dominion-manuscript/v1"

#: ``partLabel``/``volumeLabel`` join head and title with a spaced em dash (labels.ts:102,109).
LABEL_SEPARATOR = " — "

#: Whole-line scene dividers recognized by ``split_plain_markdown``.
#:
#: This deliberately matches the vocabulary the repo's own manuscript importer already accepts —
#: ``_SCENE_BREAK_RE`` in ``src/dominion/workers/memory/manuscript_split.py:34-37`` — so a file that
#: imports cleanly into the app splits into the same scenes here. That includes ``---`` and ``___``:
#: inside a chapter body, an author's horizontal rule between prose blocks IS a scene break, which is
#: how the drafts in ``book1/manuscript/`` are actually written. (``parse_blocks`` still renders a
#: stray ``---`` as a thematic rule; this only governs where a *chapter* is cut into scenes.)
_SCENE_BREAK = re.compile(
    r"^(?:(?:\*\s*){3,}|-{3,}|_{3,}|⁂|#|#{2,4}\s*scene\b.*|scene\s+[0-9]+\.?)$",
    re.IGNORECASE,
)

#: Spelled-out chapter numbers, mirroring ``manuscript_split.py:39-62`` so "Chapter Two" numbers.
_WORD_NUMBERS: dict[str, int] = {
    w: i
    for i, w in enumerate(
        "zero one two three four five six seven eight nine ten eleven twelve thirteen fourteen "
        "fifteen sixteen seventeen eighteen nineteen twenty".split()
    )
}


@dataclass
class SourceDoc:
    """One ingested document, in whichever fidelity the source could supply."""

    #: Resolved book/document title.
    title: str
    #: Populated when the source has (or can be given) book structure; ``None`` for flat sources.
    manuscript: Manuscript | None
    #: The whole document as Markdown text, for the flat "doc" mode.
    raw_text: str
    #: ``True`` when semantic markers were found (lossless), ``False`` when structure was inferred.
    structured: bool
    #: Any non-semantic YAML front matter found on the source, already stripped from ``raw_text``.
    #: The scene files in ``book1/manuscript/scenes/`` carry ``title``/``chapter``/``scene``/``pov``
    #: here, and rendering that block as body prose would be plainly wrong.
    front_matter: dict[str, str] = dc_field(default_factory=dict)


# ── Shared lexical helpers ───────────────────────────────────────────────────

#: ``<!-- chapter number=3 kind=… -->`` — the marker vocabulary emitted by ``renderMarkdown``.
_MARKER = re.compile(r"^<!--\s*(volume|part|chapter|scene)\b(.*?)-->\s*$")
#: One ``key=value`` attribute: bare, or double-quoted with ``\\`` / ``\"`` escapes (``yamlQuote``).
_ATTR = re.compile(r'([A-Za-z0-9_]+)=(?:"((?:\\.|[^"\\])*)"|([^\s"]+))')
#: A level-1 ATX heading — the only heading level ``renderMarkdown`` uses structurally.
_H1 = re.compile(r"^#\s+(.*)$")
#: Cheap probe for "this file has structure" when the front matter is missing or hand-edited.
_CHAPTER_PROBE = re.compile(r"^<!--\s*chapter\b", re.MULTILINE)
#: Fence handling mirrors ``prose.FENCE`` / ``prose.FENCE_CLOSE`` so both sides agree on extent.
_FENCE = re.compile(r"^\s*```")
_FENCE_CLOSE = re.compile(r"^\s*```\s*$")


def _normalize_newlines(text: str) -> str:
    """CRLF/CR → LF. ``renderMarkdown`` joins with ``\\n``; a CRLF copy must parse identically."""
    return text.replace("\r\n", "\n").replace("\r", "\n")


def _unescape(value: str) -> str:
    """Inverse of ``yamlQuote`` (docx.ts:1580-1582): ``\\\\`` → ``\\``, ``\\"`` → ``"``."""
    out: list[str] = []
    i = 0
    while i < len(value):
        if value[i] == "\\" and i + 1 < len(value):
            out.append(value[i + 1])
            i += 2
        else:
            out.append(value[i])
            i += 1
    return "".join(out)


def _scan_attrs(body: str) -> dict[str, str]:
    """Scan a marker comment's body into its ``key=value`` map (quoted values unescaped)."""
    attrs: dict[str, str] = {}
    for m in _ATTR.finditer(body):
        quoted, bare = m.group(2), m.group(3)
        attrs[m.group(1)] = _unescape(quoted) if quoted is not None else bare
    return attrs


def _as_int(value: str | None) -> int | None:
    """Parse an integer attribute, tolerating whitespace; ``None`` for absent/unparseable."""
    if value is None:
        return None
    try:
        return int(value.strip())
    except ValueError:
        return None


def _label_title(heading: str) -> str:
    """Recover the bare title from a rendered ``Part I — The Vault`` / ``Volume I — …`` heading.

    ``partLabel``/``volumeLabel`` build ``f"{head} — {title}"`` where ``head`` never contains the
    separator, so splitting on the FIRST occurrence is the exact inverse (labels.ts:95-110).
    """
    _, sep, title = heading.partition(LABEL_SEPARATOR)
    return title.strip() if sep else ""


def _trim_blank_edges(lines: list[str]) -> str:
    """Drop blank lines at both ends; keep internal blanks and every byte of indentation."""
    buf = list(lines)
    while buf and not buf[0].strip():
        buf.pop(0)
    while buf and not buf[-1].strip():
        buf.pop()
    return "\n".join(buf)


def _unframe_scene(lines: list[str]) -> str:
    """Undo the emitter's scene framing to recover ``prose_raw`` byte-for-byte.

    ``markdownChapter`` (docx.ts:1603-1605) pushes exactly one blank line after the scene marker
    and exactly one after the prose::

        lines.push(`<!-- scene … -->`, "");
        lines.push(sc.proseRaw);
        lines.push("");

    So precisely ONE blank line at each edge is the emitter's own framing — stripping *all* of them
    would eat a leading/trailing newline the author actually wrote. Strip one, keep the rest.
    """
    buf = list(lines)
    if buf and not buf[0].strip():
        buf.pop(0)
    if buf and not buf[-1].strip():
        buf.pop()
    return "\n".join(buf)


def _first_h1(text: str) -> str:
    """Text of the first level-1 heading, or ``""``."""
    m = re.search(r"^#\s+(.*)$", text, re.MULTILINE)
    return m.group(1).strip() if m else ""


# ── 1. Semantic Markdown: the exact inverse of renderMarkdown ────────────────


def _split_front_matter(lines: list[str]) -> tuple[dict[str, str], int]:
    """Split a leading ``---`` fenced block off the document. Returns (map, index of first body line).

    Deliberately dependency-free: ``renderMarkdown`` emits a flat ``key: value`` map with
    double-quoted strings (docx.ts:1618-1633), which is a strict subset of YAML.
    """
    if not lines or lines[0].strip() != "---":
        return {}, 0
    for j in range(1, len(lines)):
        if lines[j].strip() == "---":
            return _parse_front_matter(lines[1:j]), j + 1
    return {}, 0  # unterminated block → treat the whole document as body


def _parse_front_matter(body: list[str]) -> dict[str, str]:
    """Parse the flat ``key: value`` map. Quoted values are unquoted+unescaped, bare left verbatim."""
    meta: dict[str, str] = {}
    for raw in body:
        line = raw.strip()
        if not line or line.startswith("#"):
            continue
        key, sep, value = line.partition(":")
        if not sep:
            continue
        value = value.strip()
        if len(value) >= 2 and value.startswith('"') and value.endswith('"'):
            value = _unescape(value[1:-1])
        meta[key.strip()] = value
    return meta


def _lookahead_marker(lines: list[str], start: int) -> tuple[str, dict[str, str], int] | None:
    """After a ``# `` heading, find its structural marker comment (blank lines allowed between).

    Returns ``(kind, attrs, next_index)``, or ``None`` when the heading carries no marker — which is
    how the document's own ``# {title}`` line (docx.ts:1634) is told apart from a chapter heading.
    """
    j = start
    while j < len(lines) and not lines[j].strip():
        j += 1
    if j >= len(lines):
        return None
    m = _MARKER.match(lines[j].strip())
    if m is None or m.group(1) == "scene":
        return None
    return m.group(1), _scan_attrs(m.group(2)), j + 1


def _collect_prose(lines: list[str], start: int) -> tuple[str, int]:
    """Collect a scene's verbatim prose from ``start`` to the next marker or ``# `` heading.

    Fence-aware: ``# `` and ``<!-- … -->`` lines inside a ``` block belong to the prose, so an
    ``@interface`` panel or a shell transcript survives intact.
    """
    buf: list[str] = []
    in_fence = False
    j = start
    n = len(lines)
    while j < n:
        line = lines[j]
        if in_fence:
            if _FENCE_CLOSE.match(line):
                in_fence = False
        elif _FENCE.match(line):
            in_fence = True
        elif _H1.match(line) or _MARKER.match(line.strip()):
            break
        buf.append(line)
        j += 1
    return _unframe_scene(buf), j


def parse_semantic_markdown(text: str) -> Manuscript | None:
    """Parse semantic Markdown back into a :class:`Manuscript`, or ``None`` if it has no markers.

    Recognition: front matter ``schema: dominion-manuscript/v1`` OR any ``<!-- chapter`` marker. A
    caller that gets ``None`` should fall back to :func:`split_plain_markdown`.

    Every field is taken from the marker attributes, never re-derived from the rendered heading —
    a chapter's ``title`` comes from ``title="…"`` because the heading is ``label — title`` and the
    label is regenerated on the way out. Chapter ``position`` is the 0-based ordinal in the file, so
    ``build_spine``'s ordering reproduces reading order exactly as written.
    """
    lines = _normalize_newlines(text).split("\n")
    meta, body_start = _split_front_matter(lines)
    if meta.get("schema") != SEMANTIC_SCHEMA and not _CHAPTER_PROBE.search(text):
        return None

    ms = Manuscript(
        title=meta.get("title") or "",
        series=meta.get("series") or None,
        book_no=_as_int(meta.get("book")),
        subtitle=meta.get("subtitle") or None,
    )

    current_volume_id: str | None = None
    current_part_id: str | None = None
    current_chapter: ManuscriptChapter | None = None
    fallback_title = ""

    i = body_start
    n = len(lines)
    while i < n:
        line = lines[i]

        marker = _MARKER.match(line.strip())
        if marker is not None and marker.group(1) == "scene":
            attrs = _scan_attrs(marker.group(2))
            prose, i = _collect_prose(lines, i + 1)
            if current_chapter is not None:
                scene_no = _as_int(attrs.get("scene_no"))
                if scene_no is None:
                    scene_no = len(current_chapter.scenes) + 1
                current_chapter.scenes.append(ManuscriptScene(scene_no=scene_no, prose=prose))
            continue

        heading = _H1.match(line)
        if heading is None:
            i += 1
            continue

        found = _lookahead_marker(lines, i + 1)
        if found is None:
            # The document's own title line — no marker follows it.
            if not fallback_title:
                fallback_title = heading.group(1).strip()
            i += 1
            continue

        kind_token, attrs, i = found
        heading_text = heading.group(1).strip()

        if kind_token == "volume":
            number = _as_int(attrs.get("number")) or len(ms.volumes) + 1
            current_volume_id = f"v{number}"
            current_part_id = None
            current_chapter = None
            ms.volumes.append(
                ManuscriptVolume(
                    id=current_volume_id,
                    volume_no=number,
                    title=_label_title(heading_text),
                    subtitle=attrs.get("subtitle"),
                )
            )
        elif kind_token == "part":
            number = _as_int(attrs.get("number")) or len(ms.parts) + 1
            current_part_id = f"p{number}"
            current_chapter = None
            ms.parts.append(
                ManuscriptPart(
                    id=current_part_id,
                    part_no=number,
                    title=_label_title(heading_text),
                    volume_id=current_volume_id,
                    subtitle=attrs.get("subtitle"),
                    kind=attrs.get("kind") or "part",
                )
            )
        else:
            current_chapter = ManuscriptChapter(
                pov=attrs.get("pov", ""),
                position=len(ms.chapters),
                chapter_no=_as_int(attrs.get("number")),
                title=attrs.get("title"),
                kind=attrs.get("kind") or "chapter",
                section_type=attrs.get("section_type"),
                part_id=current_part_id,
                scenes=[],
            )
            ms.chapters.append(current_chapter)

    ms.title = ms.title or fallback_title or "Untitled"
    return ms


# ── 2. Plain Markdown: inferred structure ────────────────────────────────────

#: ``Chapter 12 — The Vault`` / ``12. The Vault`` / ``3`` → number + remainder.
_NUMBERED_HEADING = re.compile(r"^(?:chapter\s+)?([0-9]+)\b[\s:.\-—]*(.*)$", re.IGNORECASE)
#: Headings whose whole text names a numberless kind.
_KEYWORD_KINDS: frozenset[str] = frozenset({"prologue", "interlude", "epilogue"})
_NON_ALNUM = re.compile(r"[^a-z0-9]+")

# Common authored front/back-matter headings. Plain Markdown does not carry structural comments,
# so these names are the safe, explicit vocabulary we can infer without guessing from prose.
_FRONT_SECTION_TYPES: frozenset[str] = frozenset(
    {
        "copyright",
        "dedication",
        "epigraph",
        "foreword",
        "preface",
        "introduction",
        "dramatis_personae",
        "map",
        "timeline",
        "pronunciation",
    }
)
_BACK_SECTION_TYPES: frozenset[str] = frozenset(
    {
        "afterword",
        "acknowledgments",
        "appendix",
        "glossary",
        "author_note",
        "about_author",
        "author_bio",
        "preview",
    }
)


def _section_heading_map() -> dict[str, tuple[str, str]]:
    out: dict[str, tuple[str, str]] = {}
    for slug, label in SECTION_TYPES.items():
        kind = "front_matter" if slug in _FRONT_SECTION_TYPES else "back_matter"
        out[_NON_ALNUM.sub("", slug.lower())] = (kind, slug)
        out[_NON_ALNUM.sub("", label.lower())] = (kind, slug)
    # Common spelling and generic section aliases.
    out["acknowledgements"] = ("back_matter", "acknowledgments")
    out["frontmatter"] = ("front_matter", "")
    out["backmatter"] = ("back_matter", "")
    return out


_SECTION_HEADINGS = _section_heading_map()


#: ``Chapter Two — The Facility`` — spelled-out numbers are the common hand-authored form.
_WORD_NUMBERED_HEADING = re.compile(r"^chapter\s+([a-z]+)\b[\s:.\-—]*(.*)$", re.IGNORECASE)


def _classify_heading(heading: str, auto_no: int) -> tuple[str, int | None, str | None, int]:
    """Derive ``(kind, chapter_no, title, next_auto_no)`` from a plain ``# `` heading."""
    text = heading.strip()
    numbered = _NUMBERED_HEADING.match(text)
    if numbered:
        number = int(numbered.group(1))
        remainder = numbered.group(2).strip()
        return "chapter", number, (remainder or None), number + 1
    normalized = _NON_ALNUM.sub("", text.lower())
    if normalized in _KEYWORD_KINDS:
        return normalized, None, None, auto_no
    section = _SECTION_HEADINGS.get(normalized)
    if section is not None:
        kind, section_type = section
        return kind, None, (text or None), auto_no
    worded = _WORD_NUMBERED_HEADING.match(text)
    if worded and worded.group(1).lower() in _WORD_NUMBERS:
        number = _WORD_NUMBERS[worded.group(1).lower()]
        remainder = worded.group(2).strip()
        return "chapter", number, (remainder or None), number + 1
    return "chapter", auto_no, (text or None), auto_no + 1


_H2 = re.compile(r"^##\s+(.*)$")

#: A chapter heading written as PLAIN TEXT with no ``#`` — the shape of
#: book1/manuscript/chapters/realmwalkers_chapter_3_working.md, which opens with a bare
#: "Chapter Three" line followed by a bare "Reins" line. Without this the heading renders as body
#: prose and the whole file collapses into one untitled chapter.
_BARE_CHAPTER = re.compile(
    r"^(chapter\s+[A-Za-z0-9]+|prologue|interlude|epilogue)\s*$", re.IGNORECASE
)


def _looks_like_title_line(line: str) -> bool:
    """A short, unpunctuated line directly under a bare chapter heading is that chapter's title."""
    t = line.strip()
    return bool(t) and len(t) <= 60 and t[-1] not in ".!?,;:" and not t.startswith(("#", ">", "-", "|"))


def _promote_bare_headings(lines: list[str]) -> list[str]:
    """Rewrite plain-text chapter headings into Markdown so the normal splitter sees them.

    ``Chapter Three`` / ``Reins`` becomes ``# Chapter Three`` / ``## Reins``. Only a line that
    stands alone (blank line before and after, or at the very top of the file) is promoted, so a
    sentence that merely mentions a chapter is never mistaken for a heading.
    """
    out = list(lines)
    in_fence = False
    for i, line in enumerate(out):
        if _FENCE.match(line):
            in_fence = not in_fence
            continue
        if in_fence or not _BARE_CHAPTER.match(line.strip()):
            continue
        before_blank = i == 0 or not out[i - 1].strip()
        after_blank = i + 1 >= len(out) or not out[i + 1].strip()
        if not (before_blank and after_blank):
            continue
        out[i] = f"# {line.strip()}"
        # The next non-blank line, if it reads like a title, becomes the chapter subtitle.
        for j in range(i + 1, min(i + 4, len(out))):
            if not out[j].strip():
                continue
            after = j + 1 >= len(out) or not out[j + 1].strip()
            if after and _looks_like_title_line(out[j]) and not _BARE_CHAPTER.match(out[j].strip()):
                out[j] = f"## {out[j].strip()}"
            break
    return out


def _lift_subtitle(body: list[str]) -> str | None:
    """Consume a leading ``## Subtitle`` line and return it as the chapter title.

    ``# Chapter Two`` / ``## The Facility`` is how the drafts in ``book1/manuscript/chapters/`` are
    written — the level-2 line is the chapter's name, not a section inside its prose. Only a heading
    that opens the body (blank lines allowed ahead of it) is lifted.
    """
    for i, line in enumerate(body):
        if not line.strip():
            continue
        m = _H2.match(line)
        if m:
            del body[: i + 1]
            return m.group(1).strip() or None
        return None
    return None


def _split_scenes(body: list[str]) -> list[ManuscriptScene]:
    """Split a chapter body on whole-line scene dividers. Blank-only segments are dropped."""
    segments: list[list[str]] = [[]]
    in_fence = False
    for line in body:
        if in_fence:
            if _FENCE_CLOSE.match(line):
                in_fence = False
        elif _FENCE.match(line):
            in_fence = True
        elif _SCENE_BREAK.match(line.strip()):
            segments.append([])
            continue
        segments[-1].append(line)

    scenes: list[ManuscriptScene] = []
    for segment in segments:
        prose = _trim_blank_edges(segment)
        if prose:
            scenes.append(ManuscriptScene(scene_no=len(scenes) + 1, prose=prose))
    return scenes


def split_plain_markdown(
    text: str, *, title: str, front_matter: dict[str, str] | None = None
) -> Manuscript:
    """Infer a book from ordinary Markdown that carries no semantic markers.

    ``# `` headings open chapters; ``***`` / ``* * *`` / ``⁂`` / a bare ``#`` open scenes. Text before
    the first heading becomes chapter 1. With no headings at all the whole document becomes one
    chapter — and, absent any divider, one scene holding the whole text.

    ``pov`` is ``""`` for every inferred chapter: the Reader emitter suppresses the POV line when pov
    is blank (``showPov`` at docx.ts:1284), so a blank POV renders as no POV rather than an empty one.
    A chapter whose body is entirely blank keeps a single empty scene so the heading is not lost.
    """
    lines = _promote_bare_headings(_normalize_newlines(text).split("\n"))

    segments: list[tuple[str | None, list[str]]] = []
    heading: str | None = None
    buf: list[str] = []
    for line in lines:
        m = _H1.match(line)
        if m:
            segments.append((heading, buf))
            heading = m.group(1).strip()
            buf = []
        else:
            buf.append(line)
    segments.append((heading, buf))

    fm = front_matter or {}
    fm_pov = fm.get("pov", "").strip()
    fm_title = fm.get("title", "").strip() or None
    fm_chapter = _as_int(fm.get("chapter"))
    fm_scene = _as_int(fm.get("scene"))
    fm_kind = fm.get("kind", "").strip() or None
    fm_section_type = fm.get("section_type", "").strip() or None
    fm_epigraph = fm.get("epigraph", "").strip() or None

    ms = Manuscript(title=title or "Untitled")
    auto_no = 1
    for head, body in segments:
        if head is None:
            if not any(line.strip() for line in body):
                continue  # no preamble text before the first heading
            kind, chapter_no, chapter_title = "chapter", 1, None
            auto_no = 2
        else:
            kind, chapter_no, chapter_title, auto_no = _classify_heading(head, auto_no)
        subtitle = _lift_subtitle(body)
        if chapter_title is None:
            chapter_title = subtitle
        scenes = _split_scenes(body) or [ManuscriptScene(scene_no=1, prose="")]

        # A single-scene file that declares its own identity in front matter (the shape of
        # book1/manuscript/scenes/*.md) keeps that identity rather than an invented one.
        section_type = None
        epigraph = None
        if head is not None:
            inferred = _SECTION_HEADINGS.get(_NON_ALNUM.sub("", head.lower()))
            if inferred is not None:
                kind, section_type = inferred
                chapter_no = None

        if len(segments) <= 2 and fm:
            if fm_chapter is not None:
                chapter_no = fm_chapter
            if fm_title and chapter_title is None:
                chapter_title = fm_title
            if fm_scene is not None and len(scenes) == 1:
                scenes[0].scene_no = fm_scene
            if fm_kind:
                kind = fm_kind
                if kind in ("front_matter", "back_matter"):
                    chapter_no = None
            if fm_section_type:
                section_type = fm_section_type
            epigraph = fm_epigraph

        ms.chapters.append(
            ManuscriptChapter(
                pov=fm_pov,
                position=len(ms.chapters),
                chapter_no=chapter_no,
                title=chapter_title,
                kind=kind,
                section_type=section_type,
                epigraph=epigraph,
                scenes=scenes,
            )
        )
    return ms


# ── 3. DOCX: text recovery ───────────────────────────────────────────────────

#: ``Heading 1`` (style display name) or ``Heading1`` (style id, what ``DocxBuilder`` writes).
_HEADING_STYLE = re.compile(r"^Heading\s*([1-9])\b", re.IGNORECASE)
_READER_HEADING_STYLES: dict[str, int] = {
    "chapterlabel": 1,
    "rwchapterlabel": 1,
    "chapter label": 1,
    "chaptertitle": 2,
    "rwchaptertitle": 2,
    "chapter title": 2,
}


def _iter_body_blocks(document: DocxDocument) -> Iterator[Paragraph | Table]:
    """Yield paragraphs and tables IN DOCUMENT ORDER.

    ``document.paragraphs`` and ``document.tables`` are separate lists and lose the interleaving, so
    the body element's children are walked directly and dispatched on tag.
    """
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _style_names(paragraph: Paragraph) -> tuple[str, str]:
    """(display name, raw ``w:pStyle`` id) — both consulted, since a style id used but not defined in
    ``styles.xml`` makes ``paragraph.style`` fall back to Normal."""
    display = ""
    try:
        style = paragraph.style
        display = style.name or "" if style is not None else ""
    except KeyError:
        display = ""
    style_id = ""
    ppr = paragraph._p.find(qn("w:pPr"))
    if ppr is not None:
        p_style = ppr.find(qn("w:pStyle"))
        if p_style is not None:
            style_id = p_style.get(qn("w:val")) or ""
    return display, style_id


def _heading_level(paragraph: Paragraph) -> int | None:
    """Heading level for built-in headings and the formatter's chapter-label/title styles."""
    for name in _style_names(paragraph):
        normalized = name.strip().lower()
        if normalized in _READER_HEADING_STYLES:
            return _READER_HEADING_STYLES[normalized]
        m = _HEADING_STYLE.match(name.strip())
        if m:
            return int(m.group(1))
    return None


def _cell_text(cell: _Cell) -> str:
    """Cell text flattened to one line, with pipes escaped so the grid survives."""
    return " ".join(cell.text.split()).replace("|", r"\|")


def _table_to_markdown(table: Table) -> str:
    """Render a Word table as a GFM pipe table; the first row is the header."""
    rows = [[_cell_text(c) for c in row.cells] for row in table.rows]
    if not rows:
        return ""
    width = max(len(r) for r in rows)

    def line(cells: list[str]) -> str:
        return "| " + " | ".join(cells + [""] * (width - len(cells))) + " |"

    out = [line(rows[0]), "| " + " | ".join(["---"] * width) + " |"]
    out.extend(line(r) for r in rows[1:])
    return "\n".join(out)


def _cell_fill(cell: _Cell) -> str:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return ""
    shd = tc_pr.find(qn("w:shd"))
    return (shd.get(qn("w:fill")) or "").upper() if shd is not None else ""


def _table_left_border_color(table: Table) -> str:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return ""
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        return ""
    left = borders.find(qn("w:left"))
    return (left.get(qn("w:color")) or "").upper() if left is not None else ""


def _first_unique(values: list[str]) -> str:
    for value in values:
        if value.strip():
            return value.strip()
    return ""


def _label_value(text: str, label: str) -> str | None:
    compact = " ".join(text.split())
    if compact.upper().startswith(label.upper() + " "):
        return compact[len(label) :].strip()
    return None


def _status_sheet_source(table: Table) -> str | None:
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    if len(rows) < 6 or not rows or not rows[0]:
        return None
    if _first_unique(rows[0]).upper() != "CURRENT STATUS":
        return None

    fields: dict[str, str] = {}
    for row in rows:
        for raw in row:
            text = raw.strip()
            if not text:
                continue
            for label in ("Name", "Race", "Level", "Class"):
                value = _label_value(text, label)
                if value is not None:
                    fields[label] = value
            parts = [part.strip() for part in text.splitlines() if part.strip()]
            if len(parts) >= 2 and parts[0].upper() in ("HEALTH", "MANA", "STAMINA"):
                fields[parts[0].title()] = parts[1]

    abilities_index = next(
        (i for i, row in enumerate(rows) if _first_unique(row).upper() == "ABILITIES"),
        None,
    )
    if abilities_index is not None and abilities_index + 1 < len(rows):
        fields["Skills"] = _first_unique(rows[abilities_index + 1]) or "None"

    if not all(key in fields for key in ("Name", "Race", "Level", "Health", "Mana", "Stamina")):
        return None

    lines = [
        "CURRENT STATUS",
        f"Name: {fields['Name']}",
        f"Race: {fields['Race']}",
        f"Level: {fields['Level']}",
        f"Class: {fields.get('Class', 'Unassigned')}",
        "",
        f"Health: {fields['Health']}",
        f"Mana: {fields['Mana']}",
        f"Stamina: {fields['Stamina']}",
        "",
        f"Skills: {fields.get('Skills', 'None')}",
    ]
    return "```stat\n" + "\n".join(lines) + "\n```"


_COMPACT_SCAN = re.compile(
    r"^(?P<name>.+?)\s+LVL\s+(?P<level>.+?)\s+HP\s+(?P<health>.+?)"
    r"\s+STAMINA\s+(?P<stamina>.+?)\s+MANA\s+(?P<mana>.+?)"
    r"\s+RACE\s+(?P<race>.+)$",
    re.IGNORECASE,
)
_RESOURCE_READOUT = re.compile(r"^(HEALTH|MANA|STAMINA)\s+(.+)$", re.IGNORECASE)


def _readout_source(table: Table) -> str | None:
    if len(table.rows) != 1 or len(table.rows[0].cells) != 1:
        return None
    cell = table.rows[0].cells[0]
    if _cell_fill(cell) != "EEF2F6" and _table_left_border_color(table) != "46546E":
        return None

    paragraphs = [" ".join(p.text.split()) for p in cell.paragraphs if p.text.strip()]
    if not paragraphs:
        return None

    if paragraphs[0].upper() == "INTERFACE":
        body: list[str] = ["[ INTERFACE ]"]
        for line in paragraphs[1:]:
            cost = re.match(r"^COST\s+(.+)$", line, re.IGNORECASE)
            body.append(f"Cost: {cost.group(1)}" if cost else line)
        return "```\n" + "\n".join(body) + "\n```"

    joined = " ".join(paragraphs)
    scan = _COMPACT_SCAN.match(joined)
    if scan:
        fields = scan.groupdict()
        lines = [
            f"Name: {fields['name']}",
            f"Level: {fields['level']}",
            f"Health: {fields['health']}",
            f"Stamina: {fields['stamina']}",
            f"Mana: {fields['mana']}",
            f"Race: {fields['race']}",
        ]
        return "```\n" + "\n".join(lines) + "\n```"

    resource = _RESOURCE_READOUT.match(joined)
    if resource:
        return f"```\n{resource.group(1).title()}: {resource.group(2)}\n```"

    return "```\n" + "\n".join(paragraphs) + "\n```"


def _table_to_source(table: Table) -> str:
    return _status_sheet_source(table) or _readout_source(table) or _table_to_markdown(table)


def _header_has_working_draft(document: DocxDocument) -> bool:
    target = "WORKING DRAFT — PROVISIONAL"
    for section in document.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            for paragraph in header.paragraphs:
                if target in " ".join(paragraph.text.upper().split()):
                    return True
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if target in " ".join(cell.text.upper().split()):
                            return True
    return False


def _restore_working_draft_marker(blocks: list[str]) -> None:
    target = "WORKING DRAFT — PROVISIONAL"
    if any(target in " ".join(block.upper().split()) for block in blocks):
        return

    chapter_re = re.compile(r"^#\s+(?:CHAPTER\b|PROLOGUE\b|EPILOGUE\b|INTERLUDE\b)", re.IGNORECASE)
    chapter_index = next((i for i, block in enumerate(blocks) if chapter_re.match(block.strip())), None)
    if chapter_index is None:
        blocks.insert(0, target)
        return

    insert_at = chapter_index + 1
    if insert_at < len(blocks) and blocks[insert_at].lstrip().startswith("## "):
        insert_at += 1
    blocks.insert(insert_at, target)


def read_docx(path: Path) -> SourceDoc:
    """Recover Markdown text from a .docx.

    Structure is NOT inferred here — the caller decides whether to run
    :func:`split_plain_markdown` on ``raw_text``. See the module docstring for what a .docx can and
    cannot give back.
    """
    document = _open_document(str(path))
    blocks: list[str] = []
    for item in _iter_body_blocks(document):
        if isinstance(item, Table):
            rendered = _table_to_source(item)
            if rendered:
                blocks.append(rendered)
            continue
        level = _heading_level(item)
        text = item.text
        if level is not None and text.strip():
            blocks.append("#" * level + " " + text.strip())
        else:
            blocks.append(text)

    while blocks and not blocks[0].strip():
        blocks.pop(0)
    while blocks and not blocks[-1].strip():
        blocks.pop()

    if _header_has_working_draft(document):
        _restore_working_draft_marker(blocks)

    doc_title = (document.core_properties.title or "").strip()
    return SourceDoc(
        title=doc_title or path.stem,
        manuscript=None,
        raw_text="\n\n".join(blocks),
        structured=False,
    )


# ── 4. Entry points ──────────────────────────────────────────────────────────


def read_markdown(text: str) -> SourceDoc:
    """Ingest Markdown text: semantic parse when the markers are there, flat text when they are not."""
    ms = parse_semantic_markdown(text)
    if ms is not None:
        return SourceDoc(title=ms.title, manuscript=ms, raw_text=text, structured=True)

    # Non-semantic front matter is metadata, not prose. Strip it off the body so it is never
    # rendered, and hand it back so the caller can harvest title / chapter / scene / pov.
    lines = _normalize_newlines(text).split("\n")
    fm, start = _split_front_matter(lines)
    body = "\n".join(lines[start:]) if fm else text
    # Look for the title in the PROMOTED form, so a plain-text "Chapter Three" heading is found
    # here too and the document does not fall back to naming itself after the file.
    promoted = "\n".join(_promote_bare_headings(body.split("\n")))
    title = (fm.get("title", "").strip() if fm else "") or _first_h1(promoted)
    return SourceDoc(
        title=title, manuscript=None, raw_text=body, structured=False, front_matter=fm
    )


def load_source(path: Path) -> SourceDoc:
    """Load any supported source file. Falls back to the file stem when the source names no title."""
    suffix = path.suffix.lower()
    if suffix == ".doc":
        raise ValueError(
            f"{path.name}: legacy binary .doc is not supported — re-save it as .docx or .md."
        )
    if suffix == ".docx":
        doc = read_docx(path)
    else:
        # utf-8-sig so a BOM-prefixed export does not poison the front-matter fence.
        doc = read_markdown(path.read_text(encoding="utf-8-sig", errors="replace"))
    if not doc.title:
        doc.title = path.stem
    return doc
